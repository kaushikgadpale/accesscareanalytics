import streamlit as st
import pandas as pd
import io
import base64
import docx
import os
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from PIL import Image

def get_sow_template():
    """Return the default SOW template with placeholders"""
    return """
{{ServiceProviderName}}
{{ServiceProviderAddress1}}
{{ServiceProviderAddress2}}
{{ServiceProviderCity}}, {{ServiceProviderState}} {{ServiceProviderPostalCode}}

Statement of Work
{{SOWDate}}

Shape

---

# 1. PARTIES
This Statement of Work ("SOW") is entered into between:
Service Provider:
{{ServiceProviderName}}
{{ServiceProviderAddress1}}
{{ServiceProviderCity}}, {{ServiceProviderState}} {{ServiceProviderPostalCode}}

Client:
{{ClientCompanyName}}
{{ClientAddress1}}
{{ClientCity}}, {{ClientState}} {{ClientPostalCode}}

---

# 2. INTRODUCTION
This Statement of Work, effective {{SOWEffectiveDate}}, is entered into between {{ServiceProviderName}} and {{ClientCompanyName}} and shall remain in effect until completion of services or termination by either party in accordance with the terms herein.

Access Care proposes dedicated mobile dental care services to be made available for {{ClientCompanyName}} employees across selected locations.

---

# 3. ENGAGEMENT DETAILS
- Project Name: {{ProjectName}}
- SOW / Quote #: {{SOWQuoteNumber}}
- Scheduled Planning Start Date: {{ScheduledPlanningStartDate}}
- Launch of Dental Service Tour: {{ScheduledPlanningStartDate}} ➔ {{ScheduledEndDate}}
- Scheduled End Date: {{ScheduledEndDate}}
- Progressive Business Manager: {{ClientBusinessManager}}
- Access Care Client Success Manager: {{AccessCareClientSuccessManager}}
- Access Care Project Manager: {{AccessCareProjectManager}}
- Insurance Partner: {{InsurancePartner}}

---

# 4. DESCRIPTION & SCOPE OF WORK
Access Care Health LLC will deliver on-site comprehensive dental services to {{ClientCompanyName}} employees through a fully equipped mobile dental unit staffed by licensed dental professionals.

## 4.1 General Services
- Comprehensive Oral Examination
- Professional Dental Cleaning
- Digital Dental X-rays
- Oral Health Education
- Personalized Treatment Recommendations

## 4.2 Staffing
Onsite Leaders:
- Dentist
- Hygienist
- RDA's
- Trucking
- Logistics
- Practice Management / Client Success / Marketing

---

# 5. TIMELINE & DELIVERABLES
## 5.1 Project Milestones
- Planning & Scheduling: {{PlanningSchedulingTimeline}}
- On-site Dental Services: {{OnsiteDentalTimeline}}
- Utilization Report Delivery: {{UtilizationReportDate}}
- Final Feedback & Debriefing: {{FinalFeedbackDebriefDate}}

---

# 6. PRICING STRUCTURE
- Daily Service Minimum: {{DailyServiceMinimumAmount}}
- Per Chair Patient Minimum: {{PerChairPatientMinimum}}
- Chairs Per Pod: {{ChairsPerPod}}
- Claim Billing Rate: {{ClaimBillingRate}}
- Travel and Logistics: {{TravelAndLogisticsValue}}
- Claims Collection Method: {{ClaimsCollectionMethod}}

---

# 7. PROMOTIONAL & MARKETING MATERIALS
- Email Blast: Linked to Site Codes
- Digital Flyer: Linked to Site Codes
- Digital Banner: Linked to Site Codes

---

# 8. COMPLIANCE & CONFIDENTIALITY
Access Care Health LLC shall comply with all applicable local, state, and federal healthcare regulations while delivering services. All personal health information will be protected and only shared with the individual employee. Access Care will maintain strict compliance with HIPAA, GDPR, and other relevant data privacy laws throughout the duration of this engagement.

---

# 9. COMPANY ADDRESSES
Access Care Health, LLC  
3525 Ridge Meadow Pkwy  
Memphis, TN 38115  
USA

---

# 10. SIGNATURES
Access Care Health, LLC  
Signed By: {{AccessCareSignatoryName}}  
Title: {{AccessCareSignatoryTitle}}  
Date: {{AccessCareSignatureDate}}

{{ClientCompanyName}}  
Signed By: {{ClientSignatoryName}}  
Title: {{ClientSignatoryTitle}}  
Date: {{ClientSignatureDate}}
"""

def get_placeholder_fields():
    """Return a dictionary of all placeholder fields with empty values"""
    return {
        "ServiceProviderName": "Access Care Health, LLC",
        "ServiceProviderAddress1": "3525 Ridge Meadow Pkwy",
        "ServiceProviderAddress2": "",
        "ServiceProviderCity": "Memphis",
        "ServiceProviderState": "TN",
        "ServiceProviderPostalCode": "38115",
        
        "ClientCompanyName": "",
        "ClientAddress1": "",
        "ClientCity": "",
        "ClientState": "",
        "ClientPostalCode": "",
        
        "SOWDate": datetime.now().strftime("%m/%d/%Y"),
        "SOWEffectiveDate": datetime.now().strftime("%m/%d/%Y"),
        "SOWDuration": "12 months",
        
        "ProjectName": "",
        "SOWQuoteNumber": "",
        "ScheduledPlanningStartDate": "",
        "ScheduledEndDate": "",
        
        "ClientBusinessManager": "",
        "AccessCareClientSuccessManager": "",
        "AccessCareProjectManager": "",
        "InsurancePartner": "",
        
        "AccessCareSignatoryName": "",
        "AccessCareSignatoryTitle": "",
        "AccessCareSignatureDate": datetime.now().strftime("%m/%d/%Y"),
        
        "ClientSignatoryName": "",
        "ClientSignatoryTitle": "",
        "ClientSignatureDate": datetime.now().strftime("%m/%d/%Y"),
        
        "PlanningSchedulingTimeline": "",
        "OnsiteDentalTimeline": "",
        "UtilizationReportDate": "",
        "FinalFeedbackDebriefDate": "TBD",
        
        "DailyServiceMinimumAmount": "",
        "PerChairPatientMinimum": "",
        "ChairsPerPod": "",
        "ClaimBillingRate": "",
        "TravelAndLogisticsValue": "",
        "ClaimsCollectionMethod": ""
    }

def add_page_number(doc):
    """Add page numbers to the footer of the document"""
    sections = doc.sections
    for section in sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        
        run = paragraph.add_run(' of ')
        
        run = paragraph.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        
        instrText2 = OxmlElement('w:instrText')
        instrText2.set(qn('xml:space'), 'preserve')
        instrText2.text = "NUMPAGES"
        
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        
        run._element.append(fldChar3)
        run._element.append(instrText2)
        run._element.append(fldChar4)

def add_header_with_logos(doc, client_logo_path=None):
    """Add header with Access Care logo and optional client logo"""
    for section in doc.sections:
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        
        # Create a table for the header (2 columns)
        header_table = header.add_table(1, 2, width=Inches(8.5))
        
        # Access Care Logo (left)
        left_cell = header_table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        left_run = left_para.add_run()
        
        # Add Access Care logo (big_logo.png)
        try:
            left_run.add_picture('big_logo.png', width=Inches(2.5))
        except:
            # If logo file isn't found, just add text instead
            left_run = left_para.add_run("Access Care Health, LLC")
            left_run.font.size = Pt(14)
            left_run.font.bold = True
        
        # Client Logo (right) if provided
        right_cell = header_table.cell(0, 1)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if client_logo_path:
            try:
                right_run = right_para.add_run()
                right_run.add_picture(client_logo_path, width=Inches(2))
            except Exception as e:
                # If there's an error with the client logo, just skip it
                print(f"Error adding client logo: {e}")
        
        # Set spacing after the header
        header_para.space_after = Pt(12)

def create_docx(template_text, placeholder_values, client_logo_path=None):
    """Create a Word document from the template and placeholder values"""
    # Replace placeholders in the template
    for placeholder, value in placeholder_values.items():
        template_text = template_text.replace(f"{{{{{placeholder}}}}}", str(value) if value else "")
    
    # Create a new Word document
    doc = docx.Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add header with logos
    add_header_with_logos(doc, client_logo_path)
    
    # Add page numbers to footer
    add_page_number(doc)
    
    # Process each line of the template
    lines = template_text.split('\n')
    
    for line in lines:
        if line.startswith('# '):  # Main heading
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('## '):  # Subheading
            heading = doc.add_heading(line[3:], level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('- '):  # Bullet point
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line == '---':  # Horizontal line - add a thin line
            p = doc.add_paragraph()
            p.add_run('_' * 80)
        else:
            if line.strip():  # Regular paragraph
                p = doc.add_paragraph(line)
            else:  # Empty line
                doc.add_paragraph()
    
    # Save to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream

def get_docx_download_link(docx_file, filename="Statement_of_Work.docx"):
    """Generate a download link for the Word document"""
    b64 = base64.b64encode(docx_file.read()).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Word Document</a>'
    return href

def render_sow_creator():
    """Render the SOW creator interface"""
    st.header("Statement of Work Generator")
    
    # Instructions
    st.markdown("""
    <div class="card">
        <h3>SOW Creator Tool</h3>
        <p>Fill in the fields below to generate a Statement of Work document. Required fields are marked with *.</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Get the template and default values
    template = get_sow_template()
    default_values = get_placeholder_fields()
    
    # If values exist in session state, use those
    if 'sow_values' not in st.session_state:
        st.session_state.sow_values = default_values.copy()
    
    # Initialize client logo in session state if it doesn't exist
    if 'client_logo' not in st.session_state:
        st.session_state.client_logo = None
    
    # Create a tabbed interface for categories of fields
    tabs = ["Client Information", "Project Details", "Pricing & Timeline", "Signatures", "Document Settings"]
    
    tab1, tab2, tab3, tab4, tab5 = st.tabs(tabs)
    
    with tab1:
        st.subheader("Client Information")
        
        # Client Logo Upload
        uploaded_file = st.file_uploader("Upload Client Logo (optional)", type=["png", "jpg", "jpeg"])
        if uploaded_file is not None:
            # Save the uploaded file
            client_logo_path = f"client_logo_{uploaded_file.name}"
            with open(client_logo_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            # Display the uploaded logo
            st.image(uploaded_file, caption='Client Logo', width=200)
            
            # Store the path in session state
            st.session_state.client_logo = client_logo_path
        
        col1, col2 = st.columns(2)
        with col1:
            st.session_state.sow_values["ClientCompanyName"] = st.text_input(
                "Client Company Name *", 
                value=st.session_state.sow_values["ClientCompanyName"]
            )
            
            st.session_state.sow_values["ClientAddress1"] = st.text_input(
                "Client Address Line 1 *", 
                value=st.session_state.sow_values["ClientAddress1"]
            )
            
        with col2:
            st.session_state.sow_values["ClientCity"] = st.text_input(
                "Client City *", 
                value=st.session_state.sow_values["ClientCity"]
            )
            
            col2a, col2b = st.columns(2)
            with col2a:
                st.session_state.sow_values["ClientState"] = st.text_input(
                    "Client State *", 
                    value=st.session_state.sow_values["ClientState"]
                )
            with col2b:
                st.session_state.sow_values["ClientPostalCode"] = st.text_input(
                    "Client Postal Code *", 
                    value=st.session_state.sow_values["ClientPostalCode"]
                )
    
    with tab2:
        st.subheader("Project Details")
        
        col1, col2 = st.columns(2)
        with col1:
            st.session_state.sow_values["ProjectName"] = st.text_input(
                "Project Name *", 
                value=st.session_state.sow_values["ProjectName"]
            )
            
            st.session_state.sow_values["SOWQuoteNumber"] = st.text_input(
                "SOW/Quote Number *", 
                value=st.session_state.sow_values["SOWQuoteNumber"]
            )
            
            st.session_state.sow_values["ClientBusinessManager"] = st.text_input(
                "Client Business Manager", 
                value=st.session_state.sow_values["ClientBusinessManager"]
            )
        
        with col2:
            st.session_state.sow_values["ScheduledPlanningStartDate"] = st.date_input(
                "Scheduled Planning Start Date *",
                format="MM/DD/YYYY"
            ).strftime("%m/%d/%Y")
            
            st.session_state.sow_values["ScheduledEndDate"] = st.date_input(
                "Scheduled End Date *",
                format="MM/DD/YYYY"
            ).strftime("%m/%d/%Y")
            
            st.session_state.sow_values["InsurancePartner"] = st.text_input(
                "Insurance Partner", 
                value=st.session_state.sow_values["InsurancePartner"]
            )
        
        col3, col4 = st.columns(2)
        with col3:
            st.session_state.sow_values["AccessCareClientSuccessManager"] = st.text_input(
                "Access Care Client Success Manager", 
                value=st.session_state.sow_values["AccessCareClientSuccessManager"]
            )
        
        with col4:
            st.session_state.sow_values["AccessCareProjectManager"] = st.text_input(
                "Access Care Project Manager", 
                value=st.session_state.sow_values["AccessCareProjectManager"]
            )
    
    with tab3:
        st.subheader("Pricing & Timeline")
        
        col1, col2 = st.columns(2)
        with col1:
            st.session_state.sow_values["DailyServiceMinimumAmount"] = st.text_input(
                "Daily Service Minimum", 
                value=st.session_state.sow_values["DailyServiceMinimumAmount"]
            )
            
            st.session_state.sow_values["PerChairPatientMinimum"] = st.text_input(
                "Per Chair Patient Minimum", 
                value=st.session_state.sow_values["PerChairPatientMinimum"]
            )
            
            st.session_state.sow_values["ChairsPerPod"] = st.text_input(
                "Chairs Per Pod", 
                value=st.session_state.sow_values["ChairsPerPod"]
            )
        
        with col2:
            st.session_state.sow_values["ClaimBillingRate"] = st.text_input(
                "Claim Billing Rate", 
                value=st.session_state.sow_values["ClaimBillingRate"]
            )
            
            st.session_state.sow_values["TravelAndLogisticsValue"] = st.text_input(
                "Travel and Logistics", 
                value=st.session_state.sow_values["TravelAndLogisticsValue"]
            )
            
            st.session_state.sow_values["ClaimsCollectionMethod"] = st.text_input(
                "Claims Collection Method", 
                value=st.session_state.sow_values["ClaimsCollectionMethod"]
            )
        
        st.subheader("Timeline Milestones")
        
        col3, col4 = st.columns(2)
        with col3:
            st.session_state.sow_values["PlanningSchedulingTimeline"] = st.text_input(
                "Planning & Scheduling Timeline", 
                value=st.session_state.sow_values["PlanningSchedulingTimeline"] or st.session_state.sow_values["ScheduledPlanningStartDate"]
            )
            
            st.session_state.sow_values["OnsiteDentalTimeline"] = st.text_input(
                "On-site Dental Services Timeline", 
                value=st.session_state.sow_values["OnsiteDentalTimeline"] or f"{st.session_state.sow_values['ScheduledPlanningStartDate']} - {st.session_state.sow_values['ScheduledEndDate']}"
            )
        
        with col4:
            st.session_state.sow_values["UtilizationReportDate"] = st.text_input(
                "Utilization Report Delivery Date", 
                value=st.session_state.sow_values["UtilizationReportDate"]
            )
            
            st.session_state.sow_values["FinalFeedbackDebriefDate"] = st.text_input(
                "Final Feedback & Debriefing Date", 
                value=st.session_state.sow_values["FinalFeedbackDebriefDate"]
            )
    
    with tab4:
        st.subheader("Signatures")
        
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("#### Access Care")
            st.session_state.sow_values["AccessCareSignatoryName"] = st.text_input(
                "Access Care Signatory Name", 
                value=st.session_state.sow_values["AccessCareSignatoryName"]
            )
            
            st.session_state.sow_values["AccessCareSignatoryTitle"] = st.text_input(
                "Access Care Signatory Title", 
                value=st.session_state.sow_values["AccessCareSignatoryTitle"]
            )
            
            st.session_state.sow_values["AccessCareSignatureDate"] = st.date_input(
                "Access Care Signature Date",
                value=datetime.now(),
                format="MM/DD/YYYY"
            ).strftime("%m/%d/%Y")
        
        with col2:
            st.markdown("#### Client")
            st.session_state.sow_values["ClientSignatoryName"] = st.text_input(
                "Client Signatory Name", 
                value=st.session_state.sow_values["ClientSignatoryName"]
            )
            
            st.session_state.sow_values["ClientSignatoryTitle"] = st.text_input(
                "Client Signatory Title", 
                value=st.session_state.sow_values["ClientSignatoryTitle"]
            )
            
            st.session_state.sow_values["ClientSignatureDate"] = st.date_input(
                "Client Signature Date",
                value=datetime.now(),
                format="MM/DD/YYYY"
            ).strftime("%m/%d/%Y")
    
    with tab5:
        st.subheader("Document Settings")
        st.markdown("#### Layout and Formatting")
        
        # Informational text about the document appearance
        st.info("""
        The generated SOW document will include:
        - A professional header with Access Care logo and client logo (if provided)
        - Page numbers in the footer (Page X of Y format)
        - Properly formatted sections, headings, and content
        - Standard legal document styling
        """)
        
        # Display sample header/footer if available
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### Header Preview:")
            st.image("big_logo.png", width=200, caption="Access Care Logo in header")
        
        with col2:
            if st.session_state.client_logo:
                st.markdown("#### Client Logo Preview:")
                st.image(st.session_state.client_logo, width=200, caption="Client Logo in header")
    
    # Preview and Generate buttons
    st.markdown("---")
    preview_col, gen_col = st.columns([1, 1])
    
    with preview_col:
        # Check if required fields are filled
        required_fields = ["ClientCompanyName", "ProjectName", "SOWQuoteNumber", 
                          "ScheduledPlanningStartDate", "ScheduledEndDate"]
        
        missing_fields = [field for field in required_fields if not st.session_state.sow_values.get(field)]
        
        if missing_fields:
            st.warning(f"Please fill in the required fields to preview: {', '.join(missing_fields)}")
            preview_disabled = True
        else:
            preview_disabled = False
        
        if st.button("Preview SOW", disabled=preview_disabled):
            # Generate preview
            filled_template = template
            for key, value in st.session_state.sow_values.items():
                filled_template = filled_template.replace(f"{{{{{key}}}}}", str(value) if value else "")
            
            # Show in an expander
            with st.expander("SOW Preview (Markdown Format)", expanded=True):
                st.markdown(filled_template)
    
    with gen_col:
        if st.button("Generate Word Document", disabled=preview_disabled):
            # Generate DOCX with client logo if available
            client_logo_path = st.session_state.client_logo if st.session_state.client_logo else None
            docx_file = create_docx(template, st.session_state.sow_values, client_logo_path)
            
            # Provide download link
            st.markdown(get_docx_download_link(docx_file), unsafe_allow_html=True)
            
            # Also show success message
            st.success("DOCX file generated successfully. Click the link above to download.")
            
            # Show information about what was included
            st.info("The document includes headers with logos, footers with page numbers, and all the information you provided.") 